"""One-off builder for documents/summarization_prompt_reasoning.docx"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

OUT = Path(__file__).resolve().parents[1] / "documents" / "summarization_prompt_resoning.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def main() -> None:
    doc = Document()
    title = doc.add_heading("agentFlow Summarization Tool — Prompt, Context, Constraints & Reasoning", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_body(
        doc,
        "This document explains how the summarize tool in backend/tools.py works: what we send to "
        "the model (Ollama), why we chose that design, guardrails against hallucination, required "
        "output format, and exception handling. Implementation file: backend/tools.py.",
    )
    doc.add_paragraph()

    # 1
    add_heading(doc, "1. Role in the agentFlow pipeline", 1)
    add_bullet(doc, "Groq (agent.py) runs the planner, tool selector, and workflow-level summarizer.")
    add_bullet(doc, "Ollama (tools.py summarize) runs only when Groq selects the summarize tool for a subtask.")
    add_bullet(doc, "Input to summarize is task_description if set, otherwise task_name — passed as the text argument.")
    add_bullet(
        doc,
        "Output is stored in tasks.result JSONB as tool_output with keys: tool, success, input_length, summary, bullets.",
    )
    add_body(
        doc,
        "Important limitation today: summarize does not automatically receive output from a previous "
        "web_search task. The SOURCE TEXT is whatever text is in that task’s name/description unless "
        "you later extend the agent loop to pass prior tool_output.",
    )

    # 2
    add_heading(doc, "2. Why Ollama instead of Groq for this tool", 1)
    add_bullet(doc, "Separates concerns: Groq handles orchestration; local Ollama handles condensation of provided text.")
    add_bullet(doc, "Keeps summarization of raw source on your machine (no extra cloud call for that step).")
    add_bullet(doc, "Low temperature (default 0.1) reduces creative drift when condensing fixed source text.")
    add_bullet(doc, "Workflow-level final answer still uses Groq summarize_workflow in agent.py — different job.")

    # 3
    add_heading(doc, "3. Environment variables", 1)
    env_table = [
        ("OLLAMA_BASE_URL", "http://127.0.0.1:11434", "Ollama server URL"),
        ("OLLAMA_MODEL", "llama3.2", "Model tag (must be pulled: ollama pull <model>)"),
        ("OLLAMA_TIMEOUT_SECONDS", "120", "HTTP timeout for /api/chat"),
        ("OLLAMA_TEMPERATURE", "0.1", "Lower = stick closer to source wording"),
        ("SUMMARIZE_MAX_INPUT_CHARS", "32000", "Reject oversized source before LLM call"),
        ("SUMMARIZE_MAX_BULLETS", "6", "Max bullets after parse"),
        ("SUMMARIZE_MIN_BULLET_OVERLAP", "0.35", "Guardrail: fraction of content words that must appear in source"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Variable"
    hdr[1].text = "Default"
    hdr[2].text = "Purpose"
    for var, default, purpose in env_table:
        row = table.add_row().cells
        row[0].text = var
        row[1].text = default
        row[2].text = purpose
    doc.add_paragraph()

    # 4
    add_heading(doc, "4. Messages sent to Ollama", 1)
    add_heading(doc, "4.1 System prompt (SUMMARIZE_SYSTEM)", 2)
    add_code_block(
        doc,
        """You are a strict summarization tool for agentFlow.
You receive SOURCE TEXT only. Condense it — never add outside knowledge.

Mandatory rules:
1. Use ONLY facts, names, numbers, and claims explicitly present in SOURCE TEXT.
2. Do NOT invent, infer, assume, or use general world knowledge not in the source.
3. If the source is short or vague, write a minimal honest summary; do not pad with filler.
4. Return ONLY valid JSON (no markdown fences), exactly:
   {"summary": "1-3 sentences", "bullets": ["point 1", "point 2"]}
5. "summary": 1-3 sentences, plain language, every claim traceable to the source.
6. "bullets": 2-6 short bullets; each bullet = one point from the source only.
7. If nothing can be extracted, use "bullets": [] and state that in summary.""",
    )
    add_body(
        doc,
        "Reasoning: The system role sets the model’s job as compression, not research. Rules 1–2 block "
        "hallucination at the instruction level. Rules 4–7 lock the response shape so Python can parse "
        "and store results consistently.",
    )

    add_heading(doc, "4.2 User message template", 2)
    add_code_block(
        doc,
        """Summarize the following SOURCE TEXT. Do not add any information not present below.

--- SOURCE TEXT START ---
<task_description or task_name from agent>
--- SOURCE TEXT END ---""",
    )
    add_body(
        doc,
        "Reasoning: Clear delimiters show the model exactly what is “in bounds.” The agent passes "
        "stripped task text only — no hidden context from other tasks unless you add that later.",
    )

    add_heading(doc, "4.3 Ollama API payload (non-message fields)", 2)
    add_bullet(doc, "Endpoint: POST {OLLAMA_BASE_URL}/api/chat")
    add_bullet(doc, "stream: false — wait for full JSON response")
    add_bullet(doc, "format: JSON schema (see section 5) — Ollama constrains output structure")
    add_bullet(doc, "options.temperature: from OLLAMA_TEMPERATURE")

    # 5
    add_heading(doc, "5. Required JSON schema (Ollama format + Python validation)", 1)
    add_code_block(
        doc,
        """{
  "type": "object",
  "properties": {
    "summary": {"type": "string"},
    "bullets": {"type": "array", "items": {"type": "string"}}
  },
  "required": ["summary", "bullets"],
  "additionalProperties": false
}""",
    )
    add_body(doc, "After the model responds, Python also:")
    add_bullet(doc, "Strips optional ```json fences if the model ignored instructions")
    add_bullet(doc, "Requires summary: non-empty string after strip")
    add_bullet(doc, "Requires bullets: array; each item must be a string; empty strings dropped")
    add_bullet(doc, "Enforces at most SUMMARIZE_MAX_BULLETS bullets")

    add_heading(doc, "5.1 Success output shape (stored in DB)", 2)
    add_code_block(
        doc,
        """{
  "tool": "summarize",
  "success": true,
  "input_length": <character count of source>,
  "summary": "<1-3 sentences>",
  "bullets": ["<point 1>", "<point 2>", ...]
}""",
    )

    # 6
    add_heading(doc, "6. Anti-hallucination guardrails (design reasoning)", 1)
    add_body(doc, "We use three layers — prompt, schema, and code — because any single layer can fail.")

    add_heading(doc, "6.1 Layer 1 — Prompt constraints", 2)
    add_bullet(doc, "Source-only rule: no outside knowledge")
    add_bullet(doc, "Explicit JSON shape; no markdown")
    add_bullet(doc, "Allow empty bullets when source has nothing extractable (honest failure, not filler)")

    add_heading(doc, "6.2 Layer 2 — JSON schema via Ollama format", 2)
    add_body(
        doc,
        "Forces the model to emit an object with summary and bullets only. Reduces free-text answers "
        "that break the pipeline.",
    )

    add_heading(doc, "6.3 Layer 3 — Vocabulary overlap check (Python)", 2)
    add_body(
        doc,
        "After parsing JSON, we check that the summary and each bullet are “grounded” in the original "
        "source text using word overlap (not exact substring match, so light paraphrase can pass).",
    )
    add_bullet(doc, "Extract “significant words”: length > 2, not in a stopword list")
    add_bullet(doc, "Overlap ratio = (significant words in fragment that appear in source) / (total significant words in fragment)")
    add_bullet(doc, "Summary must meet overlap >= SUMMARIZE_MIN_BULLET_OVERLAP × 0.8 (slightly looser than bullets)")
    add_bullet(doc, "Each bullet must meet overlap >= SUMMARIZE_MIN_BULLET_OVERLAP (default 0.35)")
    add_bullet(doc, "If fragment is a substring of source (case-insensitive), it passes immediately")
    add_bullet(doc, "Bullets that fail overlap are dropped; if ALL bullets fail, the tool raises ToolExecutionError")
    add_body(
        doc,
        "Reasoning: Paraphrases often share root words with the source; invented facts often introduce "
        "new names or claims with weak overlap. This is a heuristic, not perfect — tune "
        "SUMMARIZE_MIN_BULLET_OVERLAP if you see false rejects or false accepts.",
    )

    # 7
    add_heading(doc, "7. Processing flow (step by step)", 1)
    steps = [
        "run_tool('summarize', ...) builds context = task_description or task_name",
        "summarize(text=context) strips and rejects empty input",
        "Reject if len(source) > SUMMARIZE_MAX_INPUT_CHARS",
        "POST to Ollama with system + user messages, format schema, low temperature",
        "Parse JSON; validate types and bullet count",
        "Apply overlap guardrails on summary and bullets",
        "Return success dict; agent.py sets task.status=completed and stores tool_output",
    ]
    for i, step in enumerate(steps, 1):
        add_bullet(doc, f"{i}. {step}")

    # 8
    add_heading(doc, "8. Exception handling (ToolExecutionError)", 1)
    add_body(
        doc,
        "All failures raise ToolExecutionError(tool_name='summarize'). agent.py catches this per task: "
        "task.status='failed', task.result includes success=false and error details; other tasks and "
        "the workflow still run to completion.",
    )
    errors = [
        ("Empty input", "Summarize input text is empty"),
        ("Input too long", "Exceeds SUMMARIZE_MAX_INPUT_CHARS"),
        ("OLLAMA_MODEL empty", "Not configured in .env"),
        ("ConnectError", "Ollama not running at OLLAMA_BASE_URL"),
        ("TimeoutException", "Request exceeded OLLAMA_TIMEOUT_SECONDS"),
        ("HTTPStatusError", "Ollama returned 4xx/5xx"),
        ("Invalid JSON", "Model output not parseable"),
        ("Schema violations", "Missing summary, wrong types, too many bullets"),
        ("Empty Ollama content", "message.content missing or done=false"),
        ("Guardrail: summary", "Summary overlap below threshold"),
        ("Guardrail: all bullets", "Every bullet failed overlap check"),
    ]
    err_table = doc.add_table(rows=1, cols=2)
    err_table.style = "Table Grid"
    err_table.rows[0].cells[0].text = "Condition"
    err_table.rows[0].cells[1].text = "User-facing message (summary)"
    for cond, msg in errors:
        r = err_table.add_row().cells
        r[0].text = cond
        r[1].text = msg

    # 9
    add_heading(doc, "9. Tool catalog text (what Groq sees when selecting tools)", 1)
    add_code_block(
        doc,
        """name: summarize
description: Summarize provided text into a short summary and bullet list using a local Ollama model.
Use only when the subtask is to condense or report on text already in the task description
(not for live web search — use web_search first if facts are missing).""",
    )
    add_body(
        doc,
        "Reasoning: Steers the tool selector toward summarize only when the task already contains "
        "text to condense, and toward web_search when facts must be fetched first.",
    )

    # 10
    add_heading(doc, "10. How to test locally", 1)
    add_bullet(doc, "Terminal 1: ollama serve")
    add_bullet(doc, "ollama pull llama3.2  (or your OLLAMA_MODEL)")
    add_bullet(doc, "Set OLLAMA_* in .env")
    add_bullet(doc, "POST /workflows with a goal that produces a summarize-friendly subtask")
    add_bullet(doc, "Or: python -c \"from backend.tools import summarize; print(summarize(text='Your source paragraph here.'))\"")
    add_bullet(doc, "Watch [DEBUG] blocks when AGENTFLOW_DEBUG=1")

    add_heading(doc, "11. Related code (for cross-reference)", 1)
    add_bullet(doc, "backend/tools.py — summarize(), _call_ollama_summarize(), guardrails")
    add_bullet(doc, "backend/agent.py — _select_and_run_tools(), ToolExecutionError handling")
    add_bullet(doc, ".env.example — Ollama and SUMMARIZE_* variables")

    doc.add_paragraph()
    p = doc.add_paragraph("Document generated for agentFlow. Aligns with backend/tools.py as of implementation date.")
    p.runs[0].font.size = Pt(9)
    p.runs[0].italic = True

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
